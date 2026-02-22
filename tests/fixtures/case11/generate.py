"""Generate a 3-page DOCX with headers, footers, and page numbers for case11."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

def add_page_number_field(paragraph):
    """Insert PAGE field code into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_char_end)

def add_numpages_field(paragraph):
    """Insert NUMPAGES field code into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " NUMPAGES "
    run2._element.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run3._element.append(fld_char_end)

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.5)
section.footer_distance = Inches(0.5)

section.different_first_page_header_footer = True

# First page header
first_header = section.first_page_header
first_header.is_linked_to_previous = False
p = first_header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONFIDENTIAL — Draft Report")
run.bold = True
run.font.size = Pt(14)

# Default header (pages 2+)
default_header = section.header
default_header.is_linked_to_previous = False
p = default_header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("Quarterly Report 2025")
run.font.size = Pt(10)

# First page footer
first_footer = section.first_page_footer
first_footer.is_linked_to_previous = False
p = first_footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Internal Use Only")
run.font.size = Pt(9)

# Default footer with "Page X of Y"
default_footer = section.footer
default_footer.is_linked_to_previous = False
p = default_footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Page ")
run.font.size = Pt(10)
add_page_number_field(p)
run = p.add_run(" of ")
run.font.size = Pt(10)
add_numpages_field(p)

# Body content - Page 1
h1 = doc.add_heading("Executive Summary", level=1)
doc.add_paragraph(
    "This quarterly report provides a comprehensive overview of our organizational "
    "performance during Q3 2025. The following sections detail key achievements, "
    "financial metrics, and strategic initiatives undertaken during this period."
)
doc.add_paragraph(
    "Our team has made significant progress across multiple fronts, including revenue "
    "growth, customer acquisition, and product development milestones. The data "
    "presented herein reflects our commitment to transparency and accountability."
)

h2 = doc.add_heading("Financial Highlights", level=2)
doc.add_paragraph(
    "Revenue increased by 23% year-over-year, driven primarily by expansion into new "
    "markets and the successful launch of our premium service tier. Operating margins "
    "improved to 18.5%, up from 15.2% in the previous quarter."
)
doc.add_paragraph(
    "Customer acquisition costs decreased by 12% while lifetime value increased by "
    "8%, indicating improved efficiency in our marketing and sales operations. These "
    "trends are expected to continue into the next fiscal year."
)

# More content to push onto page 2
doc.add_heading("Operational Review", level=2)
doc.add_paragraph(
    "Infrastructure investments totaling $4.2 million were completed on schedule and "
    "under budget. System uptime averaged 99.97% across all production environments, "
    "exceeding our target of 99.95%. The engineering team deployed 847 production "
    "releases during the quarter, a 34% increase from Q2."
)
doc.add_paragraph(
    "Employee satisfaction scores reached an all-time high of 4.6 out of 5.0, driven "
    "by new benefits programs and flexible work arrangements. Voluntary turnover "
    "decreased to 6.2%, well below the industry average of 13.5%."
)

doc.add_heading("Market Analysis", level=2)
doc.add_paragraph(
    "The competitive landscape continued to evolve during Q3, with several new entrants "
    "in our primary market segment. Despite increased competition, we maintained our "
    "market share at 28.3% and expanded our presence in the enterprise segment by 15%. "
    "Our brand recognition surveys indicate strong positioning among target demographics."
)
doc.add_paragraph(
    "International expansion efforts yielded promising results, with our EMEA region "
    "growing 31% and APAC growing 28%. Strategic partnerships established during the "
    "quarter are expected to accelerate growth in these regions through 2026."
)

# Page 2/3 content
doc.add_heading("Strategic Initiatives", level=1)
doc.add_paragraph(
    "Several key strategic initiatives were launched during Q3 to position the company "
    "for long-term growth and market leadership. These initiatives span technology, "
    "talent, and market development dimensions."
)
doc.add_heading("Technology Roadmap", level=2)
doc.add_paragraph(
    "The next-generation platform architecture entered beta testing with select "
    "enterprise customers. Early feedback has been overwhelmingly positive, with "
    "participants reporting 40% faster processing times and improved ease of use. "
    "General availability is targeted for Q1 2026."
)
doc.add_paragraph(
    "Our AI and machine learning capabilities were significantly enhanced through "
    "both internal development and strategic acquisitions. The integration of advanced "
    "natural language processing models into our product suite has opened new use cases "
    "and revenue streams that were previously inaccessible."
)
doc.add_heading("Talent Development", level=2)
doc.add_paragraph(
    "A comprehensive leadership development program was launched for mid-level managers, "
    "with 85 participants enrolled in the first cohort. Early assessments show "
    "measurable improvements in team performance metrics and employee engagement scores "
    "within participating departments."
)
doc.add_paragraph(
    "Technical hiring continued at pace, with 127 new engineers joining during Q3. "
    "Our revised interview process resulted in a 23% improvement in offer acceptance "
    "rates and a more diverse candidate pipeline. Diversity metrics improved across all "
    "categories, with women in technical roles increasing from 32% to 36%."
)

doc.add_heading("Risk Assessment and Mitigation", level=2)
doc.add_paragraph(
    "Key risks identified during the quarter include regulatory changes in our primary "
    "markets, potential supply chain disruptions, and cybersecurity threats. Mitigation "
    "strategies have been developed and approved by the board for each identified risk "
    "category. Our enterprise risk management framework continues to mature, with "
    "quarterly reviews ensuring alignment with evolving business conditions."
)
doc.add_paragraph(
    "The compliance team completed a comprehensive audit of all operational processes, "
    "resulting in 14 recommendations for improvement. All critical findings have been "
    "addressed, with remaining items on track for completion by end of Q4."
)

doc.add_heading("Looking Ahead", level=1)
doc.add_paragraph(
    "As we enter Q4 2025, our focus shifts to executing on the strategic priorities "
    "established during the annual planning cycle. Key objectives include achieving "
    "full-year revenue targets, completing the platform migration, and establishing "
    "market presence in three additional geographic regions."
)
doc.add_paragraph(
    "The executive team remains confident in our ability to deliver on these objectives "
    "while maintaining the operational excellence that has characterized our recent "
    "performance. We look forward to reporting continued progress in our Q4 review."
)

doc.save("tests/fixtures/case11/input.docx")
print("Generated tests/fixtures/case11/input.docx")
